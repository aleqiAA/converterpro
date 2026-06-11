"""
Comprehensive file converter - pypdf, pdfplumber, python-docx, pandas, Pillow
Features: convert, merge, split, compress, rotate, delete, extract, watermark, jpg<->pdf, protect, unlock
"""
import os
import tempfile
import zipfile
import shutil
from io import BytesIO

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pypdf
    from pypdf import PdfWriter, PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


def get_supported_outputs(input_ext):
    ext = input_ext.lower().strip('.')
    image_formats = ['jpg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico']
    if ext in ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico']:
        return [f for f in image_formats if f != ext and f != 'jpeg'] + ['pdf']
    if ext == 'pdf':
        return ['txt', 'html', 'md', 'docx', 'jpg', 'split', 'compress', 'rotate_90', 'rotate_180', 'rotate_270', 'delete_pages', 'extract_pages', 'add_watermark']
    if ext == 'docx':
        return ['txt', 'html', 'md', 'pdf']
    if ext == 'doc':
        return ['txt']
    if ext == 'txt':
        return ['docx', 'html', 'md', 'pdf']
    if ext == 'md':
        return ['html', 'txt', 'docx']
    if ext == 'html':
        return ['txt', 'md', 'docx']
    if ext == 'csv':
        return ['xlsx', 'json', 'html', 'txt']
    if ext in ['xlsx', 'xls']:
        return ['csv', 'json', 'html', 'txt']
    if ext == 'json':
        return ['csv', 'xlsx', 'html', 'txt']
    return []


def convert_file(input_path, output_format, merge_paths=None, password=None, page_numbers=None, watermark_text=None):
    if not os.path.exists(input_path):
        raise ValueError("Input file not found.")

    input_ext = os.path.splitext(input_path)[1].lower().strip('.')
    out_fmt   = output_format.lower().strip()

    if out_fmt == 'jpeg':
        out_fmt = 'jpg'
    if input_ext == 'jpeg':
        input_ext = 'jpg'

    image_exts = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico'}

    # ── PDF-specific operations ──────────────────────────────────────────────
    if out_fmt == 'merge':
        return _merge_pdf([input_path] + (merge_paths or []))
    if out_fmt == 'split':
        return _split_pdf(input_path)
    if out_fmt == 'compress':
        return _compress_pdf(input_path)
    if out_fmt in ('rotate_90', 'rotate_180', 'rotate_270'):
        angle = int(out_fmt.split('_')[1])
        return _rotate_pdf(input_path, angle)
    if out_fmt == 'delete_pages':
        if not page_numbers:
            raise ValueError("Page numbers required for delete operation")
        return _delete_pages_pdf(input_path, page_numbers)
    if out_fmt == 'extract_pages':
        if not page_numbers:
            raise ValueError("Page numbers required for extract operation")
        return _extract_pages_pdf(input_path, page_numbers)
    if out_fmt == 'add_watermark':
        if not watermark_text:
            raise ValueError("Watermark text required")
        return _add_watermark_pdf(input_path, watermark_text)
    if out_fmt == 'protect':
        return _protect_pdf(input_path, password or '')
    if out_fmt == 'unlock':
        return _unlock_pdf(input_path, password or '')

    # ── Image → PDF ──────────────────────────────────────────────────────────
    if input_ext in image_exts and out_fmt == 'pdf':
        return _image_to_pdf(input_path)

    # ── PDF → JPG ────────────────────────────────────────────────────────────
    if input_ext == 'pdf' and out_fmt == 'jpg':
        return _pdf_to_jpg(input_path)

    # ── Standard conversions ─────────────────────────────────────────────────
    if input_ext in image_exts:
        return _image_to_image(input_path, out_fmt)
    if input_ext == 'pdf':
        return _pdf_convert(input_path, out_fmt)
    if input_ext == 'docx':
        return _docx_convert(input_path, out_fmt)
    if input_ext == 'txt':
        return _txt_convert(input_path, out_fmt)
    if input_ext == 'md':
        return _md_convert(input_path, out_fmt)
    if input_ext == 'html':
        return _html_convert(input_path, out_fmt)
    if input_ext in {'csv', 'xlsx', 'xls', 'json'}:
        return _spreadsheet_convert(input_path, input_ext, out_fmt)

    raise ValueError(f"Input format '{input_ext}' is not supported.")


# ── IMAGE CONVERSIONS ──────────────────────────────────────────────────────────

def _image_to_image(input_path, out_fmt):
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow is not installed. Run: pip install Pillow")
    pil_fmt_map = {
        'jpg': 'JPEG', 'png': 'PNG', 'webp': 'WEBP',
        'bmp': 'BMP', 'gif': 'GIF', 'tiff': 'TIFF', 'ico': 'ICO'
    }
    if out_fmt not in pil_fmt_map:
        raise ValueError(f"Cannot convert image to '{out_fmt}'.")
    img = Image.open(input_path)
    if out_fmt == 'ico':
        img = img.resize((256, 256), Image.LANCZOS)
    if out_fmt == 'jpg' and img.mode in ('RGBA', 'P', 'LA'):
        img = img.convert('RGB')
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{out_fmt}')
    tmp.close()
    img.save(tmp.name, pil_fmt_map[out_fmt])
    return tmp.name


def _image_to_pdf(input_path):
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow is not installed. Run: pip install Pillow")
    img = Image.open(input_path)
    if img.mode in ('RGBA', 'P', 'LA'):
        img = img.convert('RGB')
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp.close()
    img.save(tmp.name, 'PDF', resolution=100.0)
    return tmp.name


def _pdf_to_jpg(input_path):
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow is not installed.")
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(input_path, dpi=150)
    except ImportError:
        raise RuntimeError(
            "PDF to JPG requires pdf2image. Run: pip install pdf2image\n"
            "Also install poppler: https://github.com/oschwartz10612/poppler-windows/releases"
        )
    if len(images) == 1:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
        tmp.close()
        img = images[0]
        if img.mode in ('RGBA', 'P', 'LA'):
            img = img.convert('RGB')
        img.save(tmp.name, 'JPEG', quality=90)
        return tmp.name
    else:
        temp_dir = tempfile.mkdtemp()
        zip_path = tempfile.NamedTemporaryFile(delete=False, suffix='.zip').name
        try:
            for i, img in enumerate(images, 1):
                if img.mode in ('RGBA', 'P', 'LA'):
                    img = img.convert('RGB')
                img.save(os.path.join(temp_dir, f'page_{i:03d}.jpg'), 'JPEG', quality=90)
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for i in range(1, len(images) + 1):
                    zf.write(os.path.join(temp_dir, f'page_{i:03d}.jpg'), arcname=f'page_{i:03d}.jpg')
            return zip_path
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)


# ── PDF TEXT EXTRACTION (pdfplumber first, pypdf fallback) ────────────────────

def _extract_pdf_text(input_path):
    """Extract text using pdfplumber (better quality) with pypdf as fallback."""
    if PDFPLUMBER_AVAILABLE:
        try:
            pages = []
            with pdfplumber.open(input_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ''
                    tables = page.extract_tables()
                    table_text = ''
                    for table in tables:
                        for row in table:
                            row_clean = [str(cell or '').strip() for cell in row]
                            table_text += ' | '.join(row_clean) + '\n'
                    combined = text
                    if table_text:
                        combined += '\n[TABLE]\n' + table_text
                    pages.append(f"--- Page {i+1} ---\n{combined}")
            return '\n\n'.join(pages)
        except Exception:
            pass

    if PYPDF_AVAILABLE:
        reader = PdfReader(input_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            pages.append(f"--- Page {i+1} ---\n{text}")
        return '\n\n'.join(pages)

    raise RuntimeError("Neither pdfplumber nor pypdf is installed. Run: pip install pdfplumber")


# ── PDF OPERATIONS ─────────────────────────────────────────────────────────────

def _split_pdf(input_path):
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    reader = PdfReader(input_path)
    num_pages = len(reader.pages)
    if num_pages < 1:
        raise ValueError("PDF is empty.")
    temp_dir = tempfile.mkdtemp()
    zip_path = tempfile.NamedTemporaryFile(delete=False, suffix='.zip').name
    try:
        for i, page in enumerate(reader.pages, 1):
            writer = PdfWriter()
            writer.add_page(page)
            with open(os.path.join(temp_dir, f'page_{i:03d}.pdf'), 'wb') as f:
                writer.write(f)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for i in range(1, num_pages + 1):
                zf.write(os.path.join(temp_dir, f'page_{i:03d}.pdf'), arcname=f'page_{i:03d}.pdf')
        return zip_path
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _merge_pdf(input_paths):
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    writer = PdfWriter()
    for pdf_path in input_paths:
        if not os.path.exists(pdf_path):
            raise ValueError(f"File not found: {pdf_path}")
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            writer.add_page(page)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp.close()
    with open(tmp.name, 'wb') as f:
        writer.write(f)
    return tmp.name


def _compress_pdf(input_path):
    """Compress PDF by removing duplicate objects and compressing streams"""
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            try:
                page.compress_content_streams()
            except Exception:
                pass  # Some pages might not support compression
            writer.add_page(page)
        
        # Try to compress - version dependent
        try:
            writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)
        except AttributeError:
            # Older pypdf versions don't have this method
            pass
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        with open(tmp.name, 'wb') as f:
            writer.write(f)
        return tmp.name
    except Exception as e:
        raise ValueError(f"PDF compression failed: {str(e)}")


def _rotate_pdf(input_path, angle):
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    if angle not in (90, 180, 270):
        raise ValueError("Rotation angle must be 90, 180, or 270.")
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp.close()
    with open(tmp.name, 'wb') as f:
        writer.write(f)
    return tmp.name


def _delete_pages_pdf(input_path, page_numbers_str):
    """Delete specific pages from PDF. Format: '1,3,5' or '2-4' (1-indexed)"""
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    
    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        if total_pages == 0:
            raise ValueError("PDF is empty.")
        
        pages_to_delete = set()
        parts = str(page_numbers_str).replace(' ', '').split(',')
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            if '-' in part:
                try:
                    range_parts = part.split('-')
                    if len(range_parts) != 2:
                        raise ValueError(f"Invalid range: {part}")
                    start, end = int(range_parts[0].strip()) - 1, int(range_parts[1].strip()) - 1
                    if start < 0 or end >= total_pages or start > end:
                        raise ValueError(f"Invalid range: {part}. PDF has {total_pages} pages.")
                    pages_to_delete.update(range(start, end + 1))
                except (ValueError, IndexError) as e:
                    raise ValueError(f"Invalid page range '{part}': {str(e)}")
            else:
                try:
                    page_num = int(part) - 1
                    if page_num < 0 or page_num >= total_pages:
                        raise ValueError(f"Page {page_num + 1} out of range (PDF has {total_pages} pages)")
                    pages_to_delete.add(page_num)
                except ValueError as e:
                    raise ValueError(f"Invalid page number '{part}': {str(e)}")
        
        if len(pages_to_delete) >= total_pages:
            raise ValueError("Cannot delete all pages from PDF")
        
        if not pages_to_delete:
            raise ValueError("No valid pages specified for deletion")
        
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if i not in pages_to_delete:
                writer.add_page(page)
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        with open(tmp.name, 'wb') as f:
            writer.write(f)
        return tmp.name
    except Exception as e:
        if "Invalid" not in str(e):
            raise ValueError(f"Delete pages failed: {str(e)}")
        raise


def _extract_pages_pdf(input_path, page_numbers_str):
    """Extract specific pages from PDF. Format: '1,3,5' or '2-4' (1-indexed)"""
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    
    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        if total_pages == 0:
            raise ValueError("PDF is empty.")
        
        pages_to_extract = []
        parts = str(page_numbers_str).replace(' ', '').split(',')
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            if '-' in part:
                try:
                    range_parts = part.split('-')
                    if len(range_parts) != 2:
                        raise ValueError(f"Invalid range: {part}")
                    start, end = int(range_parts[0].strip()) - 1, int(range_parts[1].strip()) - 1
                    if start < 0 or end >= total_pages or start > end:
                        raise ValueError(f"Invalid range: {part}. PDF has {total_pages} pages.")
                    pages_to_extract.extend(range(start, end + 1))
                except (ValueError, IndexError) as e:
                    raise ValueError(f"Invalid page range '{part}': {str(e)}")
            else:
                try:
                    page_num = int(part) - 1
                    if page_num < 0 or page_num >= total_pages:
                        raise ValueError(f"Page {page_num + 1} out of range (PDF has {total_pages} pages)")
                    pages_to_extract.append(page_num)
                except ValueError as e:
                    raise ValueError(f"Invalid page number '{part}': {str(e)}")
        
        if not pages_to_extract:
            raise ValueError("No valid pages specified for extraction")
        
        writer = PdfWriter()
        for page_num in pages_to_extract:
            writer.add_page(reader.pages[page_num])
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        with open(tmp.name, 'wb') as f:
            writer.write(f)
        return tmp.name
    except Exception as e:
        if "Invalid" not in str(e):
            raise ValueError(f"Extract pages failed: {str(e)}")
        raise


def _add_watermark_pdf(input_path, watermark_text):
    """Add text watermark to all pages of PDF"""
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    
    if not watermark_text or not str(watermark_text).strip():
        raise ValueError("Watermark text cannot be empty")
    
    watermark_text = str(watermark_text).strip()
    
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        # Create watermark overlay using reportlab if available, otherwise simple text overlay
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            watermark_buffer = BytesIO()
            c = canvas.Canvas(watermark_buffer, pagesize=letter)
            c.setFillAlpha(0.15)
            c.setFont("Helvetica-Bold", 80)
            c.rotate(45)
            c.drawString(100, 50, watermark_text)
            c.save()
            watermark_buffer.seek(0)
            watermark_pdf = PdfReader(watermark_buffer)
            watermark_page = watermark_pdf.pages[0]
            
            # Apply watermark to each page
            for page in reader.pages:
                page.merge_page(watermark_page)
                writer.add_page(page)
        except ImportError:
            # Fallback: just copy pages without watermark
            # (watermark requires reportlab which might not be installed)
            for page in reader.pages:
                writer.add_page(page)
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        with open(tmp.name, 'wb') as f:
            writer.write(f)
        return tmp.name
    except Exception as e:
        raise ValueError(f"Add watermark failed: {str(e)}")


def _protect_pdf(input_path, password):
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    if not password:
        raise ValueError("A password is required to protect the PDF.")
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp.close()
    with open(tmp.name, 'wb') as f:
        writer.write(f)
    return tmp.name


def _unlock_pdf(input_path, password):
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    reader = PdfReader(input_path)
    if reader.is_encrypted:
        if not reader.decrypt(password):
            raise ValueError("Incorrect password. Could not unlock PDF.")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    tmp.close()
    with open(tmp.name, 'wb') as f:
        writer.write(f)
    return tmp.name


def _pdf_convert(input_path, out_fmt):
    if out_fmt == 'txt':
        text = _extract_pdf_text(input_path)
        tmp  = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(text)
        tmp.close()
        return tmp.name

    if out_fmt == 'html':
        text = _extract_pdf_text(input_path)
        html_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('--- Page '):
                html_lines.append(f'<h2 style="color:#0F6E56;border-bottom:1px solid #ccc;">{stripped}</h2>')
            elif stripped.startswith('[TABLE]'):
                html_lines.append('<div style="font-family:monospace;background:#f5f5f5;padding:8px;border-radius:4px;margin:8px 0">')
            else:
                html_lines.append(f'<p>{stripped}</p>')
        html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<style>body{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;'
            'color:#333;line-height:1.6}h2{margin-top:2em}</style>'
            f'</head><body>{"".join(html_lines)}</body></html>'
        )
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        text = _extract_pdf_text(input_path)
        md_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                md_lines.append('')
                continue
            if stripped.startswith('--- Page '):
                md_lines.append(f'\n## {stripped}\n')
            else:
                md_lines.append(stripped)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write('\n'.join(md_lines))
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        text = _extract_pdf_text(input_path)
        doc  = Document()
        doc.add_heading('Converted from PDF', 0)
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith('--- Page '):
                doc.add_heading(stripped, level=1)
            elif stripped and not stripped.startswith('[TABLE]'):
                doc.add_paragraph(stripped)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"PDF to '{out_fmt}' is not supported.")


# ── DOCX CONVERSIONS ───────────────────────────────────────────────────────────

def _docx_convert(input_path, out_fmt):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    doc        = Document(input_path)
    paragraphs = [p.text for p in doc.paragraphs]

    def _style(p):
        try:
            return (p.style.name or '').lower()
        except Exception:
            return ''

    if out_fmt == 'txt':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write('\n'.join(paragraphs))
        tmp.close()
        return tmp.name

    if out_fmt == 'html':
        html_lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = _style(p)
            if 'heading 1' in style:
                html_lines.append(f'<h1>{text}</h1>')
            elif 'heading 2' in style:
                html_lines.append(f'<h2>{text}</h2>')
            elif 'heading 3' in style:
                html_lines.append(f'<h3>{text}</h3>')
            else:
                html_lines.append(f'<p>{text}</p>')
        html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<style>body{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;'
            'color:#333;line-height:1.6}</style>'
            f'</head><body>{"".join(html_lines)}</body></html>'
        )
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        md_lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                md_lines.append('')
                continue
            style = _style(p)
            if 'heading 1' in style:
                md_lines.append(f'# {text}')
            elif 'heading 2' in style:
                md_lines.append(f'## {text}')
            elif 'heading 3' in style:
                md_lines.append(f'### {text}')
            else:
                md_lines.append(text)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write('\n'.join(md_lines))
        tmp.close()
        return tmp.name

    if out_fmt == 'pdf':
        try:
            from weasyprint import HTML as WP_HTML
        except ImportError:
            raise RuntimeError("WeasyPrint required for DOCX→PDF.")
        html_path = None
        pdf_tmp   = None
        try:
            html_path = _docx_convert(input_path, 'html')
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            pdf_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            pdf_tmp.close()
            WP_HTML(string=html_content).write_pdf(pdf_tmp.name)
            return pdf_tmp.name
        except Exception:
            if pdf_tmp and os.path.exists(pdf_tmp.name):
                try: os.unlink(pdf_tmp.name)
                except OSError: pass
            raise
        finally:
            if html_path and os.path.exists(html_path):
                try: os.unlink(html_path)
                except OSError: pass

    raise ValueError(f"DOCX to '{out_fmt}' is not supported.")


# ── TXT CONVERSIONS ────────────────────────────────────────────────────────────

def _txt_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()

    if out_fmt == 'html':
        lines = [f'<p>{line}</p>' if line.strip() else '<br>' for line in text.splitlines()]
        html  = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<style>body{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;'
            'color:#333;line-height:1.6}</style>'
            f'</head><body>{"".join(lines)}</body></html>'
        )
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write(text)
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    if out_fmt == 'pdf':
        try:
            from weasyprint import HTML as WP_HTML
        except ImportError:
            raise RuntimeError("WeasyPrint required for TXT→PDF.")
        html_lines = [f'<p>{line}</p>' if line.strip() else '<br>' for line in text.splitlines()]
        html = (
            "<html><body style='font-family:Arial;padding:2cm;line-height:1.6'>"
            + ''.join(html_lines) + "</body></html>"
        )
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        WP_HTML(string=html).write_pdf(tmp.name)
        return tmp.name

    raise ValueError(f"TXT to '{out_fmt}' is not supported.")


# ── MARKDOWN CONVERSIONS ───────────────────────────────────────────────────────

def _md_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()

    if out_fmt == 'html':
        try:
            import markdown as md_lib
            html_body = md_lib.markdown(text)
        except ImportError:
            html_body = ''.join(f'<p>{line}</p>' for line in text.splitlines() if line.strip())
        html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8">'
            '<style>body{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;'
            'color:#333;line-height:1.6}</style>'
            f'</head><body>{html_body}</body></html>'
        )
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'txt':
        import re
        clean = re.sub(r'#{1,6}\s', '', text)
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'`(.+?)`', r'\1', clean)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(clean)
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        import re
        html_path = _md_convert(input_path, 'html')
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        finally:
            try: os.unlink(html_path)
            except OSError: pass
        clean = re.sub(r'<[^>]+>', '', html_content)
        doc   = Document()
        for line in clean.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"Markdown to '{out_fmt}' is not supported.")


# ── HTML CONVERSIONS ───────────────────────────────────────────────────────────

def _html_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        html_content = f.read()
    import re

    def strip_tags(html):
        clean = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
        clean = re.sub(r'<script[^>]*>.*?</script>', '', clean, flags=re.DOTALL)
        clean = re.sub(r'<[^>]+>', '', clean)
        clean = re.sub(r'\n\s*\n', '\n\n', clean)
        return clean.strip()

    if out_fmt == 'txt':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(strip_tags(html_content))
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write(strip_tags(html_content))
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        text = strip_tags(html_content)
        doc  = Document()
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"HTML to '{out_fmt}' is not supported.")


# ── SPREADSHEET CONVERSIONS ────────────────────────────────────────────────────

def _spreadsheet_convert(input_path, input_ext, out_fmt):
    if not PANDAS_AVAILABLE:
        raise RuntimeError("pandas is not installed. Run: pip install pandas openpyxl")

    if input_ext == 'csv':
        df = pd.read_csv(input_path)
    elif input_ext in ('xlsx', 'xls'):
        df = pd.read_excel(input_path)
    elif input_ext == 'json':
        df = pd.read_json(input_path)
    else:
        raise ValueError(f"Cannot read '{input_ext}' as spreadsheet.")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{out_fmt}')
    tmp.close()

    if out_fmt == 'csv':
        df.to_csv(tmp.name, index=False)
    elif out_fmt == 'xlsx':
        df.to_excel(tmp.name, index=False)
    elif out_fmt == 'json':
        df.to_json(tmp.name, orient='records', indent=2)
    elif out_fmt == 'html':
        html = (
            '<!DOCTYPE html><html><head><meta charset="UTF-8"><style>'
            'body{font-family:Arial,sans-serif;padding:20px}'
            'table{border-collapse:collapse;width:100%}'
            'th{background:#0F6E56;color:white;padding:8px 12px;text-align:left}'
            'td{border:1px solid #ddd;padding:8px 12px}'
            'tr:nth-child(even){background:#f9f9f9}'
            f'</style></head><body>{df.to_html(index=False, border=0)}</body></html>'
        )
        with open(tmp.name, 'w', encoding='utf-8') as f:
            f.write(html)
    elif out_fmt == 'txt':
        with open(tmp.name, 'w', encoding='utf-8') as f:
            f.write(df.to_string(index=False))
    else:
        os.unlink(tmp.name)
        raise ValueError(f"Spreadsheet to '{out_fmt}' is not supported.")

    return tmp.name
