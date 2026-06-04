"""
Comprehensive file converter - uses only: pypdf, python-docx, pandas, Pillow
No PyPDF2, no pdf2docx, no poppler needed.
"""
import os
import tempfile

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


def get_supported_outputs(input_ext):
    """Return list of supported output formats for a given input extension."""
    ext = input_ext.lower().strip('.')

    image_formats = ['jpg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico']
    doc_formats = ['txt', 'html', 'md', 'docx']
    sheet_formats = ['csv', 'xlsx', 'json']

    if ext in ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico']:
        return [f for f in image_formats if f != ext and f != 'jpeg']

    if ext == 'pdf':
        return ['txt', 'html', 'md', 'docx']

    if ext == 'docx':
        return ['txt', 'html', 'md', 'pdf']

    if ext == 'doc':
        return ['txt']

    if ext == 'txt':
        return ['docx', 'html', 'md', 'pdf']

    if ext == 'md':
        return ['html', 'txt', 'docx']

    if ext == 'html':
        return ['txt', 'md', 'docx']

    if ext in ['csv']:
        return ['xlsx', 'json', 'html', 'txt']

    if ext in ['xlsx', 'xls']:
        return ['csv', 'json', 'html', 'txt']

    if ext == 'json':
        return ['csv', 'xlsx', 'html', 'txt']

    return []


def convert_file(input_path, output_format):
    """
    Convert input_path to output_format.
    Returns path to the converted file (caller must delete it).
    Raises ValueError for unsupported conversions.
    Raises RuntimeError for conversion failures.
    """
    if not os.path.exists(input_path):
        raise ValueError("Input file not found.")

    input_ext = os.path.splitext(input_path)[1].lower().strip('.')
    out_fmt = output_format.lower().strip('.')

    # Normalize jpeg->jpg
    if out_fmt == 'jpeg':
        out_fmt = 'jpg'
    if input_ext == 'jpeg':
        input_ext = 'jpg'

    # Route to the right converter
    image_exts = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'ico'}

    if input_ext in image_exts:
        return _image_to_image(input_path, out_fmt)

    if input_ext == 'pdf':
        return _pdf_convert(input_path, out_fmt)

    if input_ext == 'docx':
        return _docx_convert(input_path, out_fmt)

    if input_ext == 'txt':
        return _txt_convert(input_path, out_fmt)

    if input_ext == 'md':
        return _md_convert(input_path, out_fmt)

    if input_ext == 'html':
        return _html_convert(input_path, out_fmt)

    if input_ext in {'csv', 'xlsx', 'xls', 'json'}:
        return _spreadsheet_convert(input_path, input_ext, out_fmt)

    raise ValueError(f"Input format '{input_ext}' is not supported.")


# ── IMAGE CONVERSIONS ──────────────────────────────────────────────────────────

def _image_to_image(input_path, out_fmt):
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow is not installed. Run: pip install Pillow")

    pil_fmt_map = {
        'jpg': 'JPEG', 'png': 'PNG', 'webp': 'WEBP',
        'bmp': 'BMP', 'gif': 'GIF', 'tiff': 'TIFF', 'ico': 'ICO'
    }

    if out_fmt not in pil_fmt_map:
        raise ValueError(f"Cannot convert image to '{out_fmt}'.")

    img = Image.open(input_path)

    # ICO size limit
    if out_fmt == 'ico':
        img = img.resize((256, 256), Image.LANCZOS)

    # JPEG needs RGB
    if out_fmt == 'jpg' and img.mode in ('RGBA', 'P', 'LA'):
        img = img.convert('RGB')

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{out_fmt}')
    tmp.close()
    img.save(tmp.name, pil_fmt_map[out_fmt])
    return tmp.name


# ── PDF CONVERSIONS ────────────────────────────────────────────────────────────

def _extract_pdf_text(input_path):
    """Extract all text from a PDF using pypdf."""
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    reader = pypdf.PdfReader(input_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ''
        pages.append(f"--- Page {i+1} ---\n{text}")
    return '\n\n'.join(pages)


def _pdf_convert(input_path, out_fmt):
    if out_fmt == 'txt':
        text = _extract_pdf_text(input_path)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(text)
        tmp.close()
        return tmp.name

    if out_fmt == 'html':
        text = _extract_pdf_text(input_path)
        html_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('--- Page '):
                html_lines.append(f'<h2 style="color:#0F6E56;border-bottom:1px solid #ccc;">{stripped}</h2>')
            else:
                html_lines.append(f'<p>{stripped}</p>')
        html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;color:#333;line-height:1.6}}h2{{margin-top:2em}}</style>
</head><body>{''.join(html_lines)}</body></html>"""
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        text = _extract_pdf_text(input_path)
        md_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                md_lines.append('')
                continue
            if stripped.startswith('--- Page '):
                md_lines.append(f'\n## {stripped}\n')
            else:
                md_lines.append(stripped)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write('\n'.join(md_lines))
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        text = _extract_pdf_text(input_path)
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith('--- Page '):
                doc.add_heading(stripped, level=1)
            elif stripped:
                doc.add_paragraph(stripped)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"PDF to '{out_fmt}' is not supported.")


# ── DOCX CONVERSIONS ───────────────────────────────────────────────────────────

def _docx_convert(input_path, out_fmt):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(input_path)
    paragraphs = [p.text for p in doc.paragraphs]

    if out_fmt == 'txt':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write('\n'.join(paragraphs))
        tmp.close()
        return tmp.name

    if out_fmt == 'html':
        html_lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name.lower()
            if 'heading 1' in style:
                html_lines.append(f'<h1>{text}</h1>')
            elif 'heading 2' in style:
                html_lines.append(f'<h2>{text}</h2>')
            elif 'heading 3' in style:
                html_lines.append(f'<h3>{text}</h3>')
            else:
                html_lines.append(f'<p>{text}</p>')
        html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;color:#333;line-height:1.6}}</style>
</head><body>{''.join(html_lines)}</body></html>"""
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        md_lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                md_lines.append('')
                continue
            style = p.style.name.lower()
            if 'heading 1' in style:
                md_lines.append(f'# {text}')
            elif 'heading 2' in style:
                md_lines.append(f'## {text}')
            elif 'heading 3' in style:
                md_lines.append(f'### {text}')
            else:
                md_lines.append(text)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write('\n'.join(md_lines))
        tmp.close()
        return tmp.name

    if out_fmt == 'pdf':
        # Render via WeasyPrint if available
        try:
            from weasyprint import HTML as WP_HTML
            html_path = _docx_convert(input_path, 'html')
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            os.unlink(html_path)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            tmp.close()
            WP_HTML(string=html_content).write_pdf(tmp.name)
            return tmp.name
        except ImportError:
            raise RuntimeError("WeasyPrint is required for DOCX to PDF. Run: pip install weasyprint")

    raise ValueError(f"DOCX to '{out_fmt}' is not supported.")


# ── TXT CONVERSIONS ────────────────────────────────────────────────────────────

def _txt_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()

    if out_fmt == 'html':
        lines = [f'<p>{line}</p>' if line.strip() else '<br>' for line in text.splitlines()]
        html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;color:#333;line-height:1.6}}</style>
</head><body>{''.join(lines)}</body></html>"""
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write(text)
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    if out_fmt == 'pdf':
        try:
            from weasyprint import HTML as WP_HTML
            html_lines = [f'<p>{line}</p>' if line.strip() else '<br>' for line in text.splitlines()]
            html = f"<html><body style='font-family:Arial;padding:2cm'>{''.join(html_lines)}</body></html>"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            tmp.close()
            WP_HTML(string=html).write_pdf(tmp.name)
            return tmp.name
        except ImportError:
            raise RuntimeError("WeasyPrint required for TXT to PDF.")

    raise ValueError(f"TXT to '{out_fmt}' is not supported.")


# ── MARKDOWN CONVERSIONS ───────────────────────────────────────────────────────

def _md_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()

    if out_fmt == 'html':
        try:
            import markdown as md_lib
            html_body = md_lib.markdown(text)
        except ImportError:
            # Fallback: basic line conversion
            html_body = ''.join(f'<p>{line}</p>' for line in text.splitlines() if line.strip())
        html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>body{{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;color:#333;line-height:1.6}}</style>
</head><body>{html_body}</body></html>"""
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
        tmp.write(html)
        tmp.close()
        return tmp.name

    if out_fmt == 'txt':
        # Strip markdown syntax crudely
        import re
        clean = re.sub(r'#{1,6}\s', '', text)
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'`(.+?)`', r'\1', clean)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(clean)
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        html_path = _md_convert(input_path, 'html')
        # Convert the HTML to DOCX
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        os.unlink(html_path)
        # Simple: strip tags and put in docx
        import re
        clean = re.sub(r'<[^>]+>', '', html_content)
        doc = Document()
        for line in clean.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"Markdown to '{out_fmt}' is not supported.")


# ── HTML CONVERSIONS ───────────────────────────────────────────────────────────

def _html_convert(input_path, out_fmt):
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        html_content = f.read()

    import re

    def strip_tags(html):
        clean = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
        clean = re.sub(r'<script[^>]*>.*?</script>', '', clean, flags=re.DOTALL)
        clean = re.sub(r'<[^>]+>', '', clean)
        clean = re.sub(r'\n\s*\n', '\n\n', clean)
        return clean.strip()

    if out_fmt == 'txt':
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        tmp.write(strip_tags(html_content))
        tmp.close()
        return tmp.name

    if out_fmt == 'md':
        text = strip_tags(html_content)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8')
        tmp.write(text)
        tmp.close()
        return tmp.name

    if out_fmt == 'docx':
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")
        text = strip_tags(html_content)
        doc = Document()
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        doc.save(tmp.name)
        return tmp.name

    raise ValueError(f"HTML to '{out_fmt}' is not supported.")


# ── SPREADSHEET CONVERSIONS ────────────────────────────────────────────────────

def _spreadsheet_convert(input_path, input_ext, out_fmt):
    if not PANDAS_AVAILABLE:
        raise RuntimeError("pandas is not installed. Run: pip install pandas openpyxl")

    # Load
    if input_ext == 'csv':
        df = pd.read_csv(input_path)
    elif input_ext in ('xlsx', 'xls'):
        df = pd.read_excel(input_path)
    elif input_ext == 'json':
        df = pd.read_json(input_path)
    else:
        raise ValueError(f"Cannot read '{input_ext}' as spreadsheet.")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{out_fmt}')
    tmp.close()

    if out_fmt == 'csv':
        df.to_csv(tmp.name, index=False)
    elif out_fmt == 'xlsx':
        df.to_excel(tmp.name, index=False)
    elif out_fmt == 'json':
        df.to_json(tmp.name, orient='records', indent=2)
    elif out_fmt == 'html':
        html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
body{{font-family:Arial,sans-serif;padding:20px}}
table{{border-collapse:collapse;width:100%}}
th{{background:#0F6E56;color:white;padding:8px 12px;text-align:left}}
td{{border:1px solid #ddd;padding:8px 12px}}
tr:nth-child(even){{background:#f9f9f9}}
</style></head><body>
{df.to_html(index=False, border=0)}
</body></html>"""
        with open(tmp.name, 'w', encoding='utf-8') as f:
            f.write(html)
    elif out_fmt == 'txt':
        with open(tmp.name, 'w', encoding='utf-8') as f:
            f.write(df.to_string(index=False))
    else:
        os.unlink(tmp.name)
        raise ValueError(f"Spreadsheet to '{out_fmt}' is not supported.")

    return tmp.name